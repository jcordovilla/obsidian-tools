#!/usr/bin/env python3
"""
Obsidian Note to DOCX Exporter

Converts a single Obsidian markdown note to DOCX using a template file.
Maps markdown elements to the template's existing styles.

Usage:
    python export_to_docx.py --note note.md --template template.docx              # dry run
    python export_to_docx.py --note note.md --template template.docx --no-dry-run # export
    python export_to_docx.py --template template.docx --list-styles               # inspect
"""

import argparse
import json
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("Error: python-docx not installed. Install with: pip install python-docx")
    sys.exit(1)

try:
    import mistune
except ImportError:
    print("Error: mistune not installed. Install with: pip install 'mistune>=3.0.0'")
    sys.exit(1)

from obsidian_utils import validate_vault_path
from docx_format import (
    remove_empty_paragraphs,
    normalize_tables,
    disable_heading_autonumber,
)


# =============================================================================
# Custom mistune plugins for Obsidian syntax
# =============================================================================

# Precompile embed regex for preprocessing (converts ![[img]] to standard ![](obsidian-embed:...))
_EMBED_RE = re.compile(r'!\[\[([^\]|]+?)(?:\|[^\]]*?)?\]\]')

# Callout prefix pattern: > [!type] or > [!type]+
_CALLOUT_RE = re.compile(r'^>\s*\[!(\w+)\]\+?\s*$', re.MULTILINE)


def _preprocess_embeds(text):
    """Convert Obsidian ![[image.png|size]] embeds to standard markdown images."""
    def _replace(m):
        target = m.group(1)
        return f'![{target}](obsidian-embed:{target})'
    return _EMBED_RE.sub(_replace, text)


def _preprocess_callouts_and_tables(text):
    """Preprocess Obsidian markdown before parsing:

    1. Strip callout prefixes (> [!type]) from blockquotes
    2. Extract tables from inside blockquotes so they parse as proper tables
    """
    lines = text.split('\n')
    result = []
    i = 0

    while i < len(lines):
        line = lines[i]

        # Detect callout prefix line: > [!type] or > [!type]+
        if _CALLOUT_RE.match(line):
            # Skip the callout prefix line entirely
            i += 1
            # Also skip a blank blockquote line (> followed by empty)
            if i < len(lines) and lines[i].strip() in ('>', '>'):
                i += 1
            continue

        # Detect table inside blockquote: lines starting with > |
        if re.match(r'^>\s*\|', line):
            # Collect all consecutive > | lines (the table)
            table_lines = []
            while i < len(lines) and re.match(r'^>\s*\|', lines[i]):
                # Strip the > prefix
                table_line = re.sub(r'^>\s?', '', lines[i])
                table_lines.append(table_line)
                i += 1
            # Insert the table outside the blockquote
            result.append('')  # blank line before table
            result.extend(table_lines)
            result.append('')  # blank line after table
            continue

        result.append(line)
        i += 1

    return '\n'.join(result)


def wikilink_plugin(md):
    """Plugin to parse [[target]] and [[target|alias]] wikilinks."""
    WIKILINK_PATTERN = r'\[\[(?P<wl_target>[^\]|]+?)(?:\|(?P<wl_alias>[^\]]+?))?\]\]'

    def parse_wikilink(inline, m, state):
        target = m.group('wl_target') or ''
        alias = m.group('wl_alias') or target
        state.append_token({
            'type': 'wikilink',
            'children': [{'type': 'text', 'raw': alias.strip()}],
            'attrs': {'target': target.strip(), 'alias': alias.strip()},
        })
        return m.end()

    md.inline.register('wikilink', WIKILINK_PATTERN, parse_wikilink, before='link')


# =============================================================================
# Main exporter class
# =============================================================================

class MarkdownToDocxExporter:
    """Export a single Obsidian markdown note to DOCX using a template."""

    DEFAULT_STYLE_MAP = {
        'title': 'Title',
        'heading_1': 'Heading 1',
        'heading_2': 'Heading 2',
        'heading_3': 'Heading 3',
        'heading_4': 'Heading 4',
        'heading_5': 'Heading 5',
        'heading_6': 'Heading 6',
        'paragraph': 'Normal',
        'code_block': 'Quote',
        'block_quote': 'Quote',
        'list_bullet': 'List Bullet',
        'list_bullet_1': 'List Bullet 2',
        'list_bullet_2': 'List Bullet 3',
        'list_number': 'List Number',
        'list_number_1': 'List Number 2',
        'list_number_2': 'List Number 3',
        'table': 'Table Grid',
    }

    IMAGE_EXTENSIONS = {'.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff', '.tif'}

    def __init__(self, note_path, template_path, output_path, vault_path,
                 dry_run=True, style_map_overrides=None,
                 no_title_promotion=False, sanitize=True,
                 keep_heading_numbers=False):
        self.note_path = Path(note_path) if note_path else None
        self.template_path = Path(template_path)
        self.output_path = Path(output_path) if output_path else None
        self.vault_path = Path(vault_path) if vault_path else None
        self.dry_run = dry_run
        self.doc = None
        self.no_title_promotion = no_title_promotion
        # Clean-by-default: drop empty paragraphs and give tables a full-width
        # layout. Prose exports also drop inherited heading auto-numbering
        # (the template's numbered headings are wrong for free-form prose),
        # unless the caller opts to keep them.
        self.sanitize = sanitize
        self.keep_heading_numbers = keep_heading_numbers
        self._title_rendered = False  # tracks whether first H1 has been rendered

        self.style_map = dict(self.DEFAULT_STYLE_MAP)
        if style_map_overrides:
            self.style_map.update(style_map_overrides)

        # Stats for dry-run reporting
        self.stats = {
            'headings': {},
            'paragraphs': 0,
            'lists': {'bullet': 0, 'numbered': 0},
            'code_blocks': 0,
            'block_quotes': 0,
            'tables': 0,
            'images': {'found': 0, 'missing': []},
            'wikilinks': 0,
            'hyperlinks': 0,
        }

        # Set up mistune parser with AST output (renderer=None) and custom plugins
        self.md = mistune.create_markdown(
            renderer=None,
            plugins=[
                'table',
                'strikethrough',
                wikilink_plugin,
            ],
        )

    # -------------------------------------------------------------------------
    # Template style discovery
    # -------------------------------------------------------------------------

    def list_template_styles(self):
        """Discover and print all styles available in the template."""
        doc = Document(str(self.template_path))
        categories = {
            'PARAGRAPH': [],
            'CHARACTER': [],
            'TABLE': [],
            'LIST': [],
        }

        type_map = {
            WD_STYLE_TYPE.PARAGRAPH: 'PARAGRAPH',
            WD_STYLE_TYPE.CHARACTER: 'CHARACTER',
            WD_STYLE_TYPE.TABLE: 'TABLE',
            WD_STYLE_TYPE.LIST: 'LIST',
        }

        for style in doc.styles:
            if style.hidden:
                continue
            category = type_map.get(style.type)
            if category:
                builtin_tag = ' (built-in)' if style.builtin else ''
                categories[category].append(f'{style.name}{builtin_tag}')

        print(f'Template: {self.template_path}')
        print('=' * 60)
        for cat_name, styles in categories.items():
            if styles:
                print(f'\n{cat_name} STYLES ({len(styles)}):')
                for s in sorted(styles):
                    print(f'  - {s}')

        print(f'\nDefault style mapping (override with --style-map JSON):')
        print(json.dumps(self.DEFAULT_STYLE_MAP, indent=2))

    # -------------------------------------------------------------------------
    # Frontmatter stripping
    # -------------------------------------------------------------------------

    @staticmethod
    def strip_frontmatter(text):
        """Remove YAML frontmatter delimited by --- ... --- from the start."""
        if text.startswith('---'):
            match = re.match(r'^---\n.*?\n---\n?', text, re.DOTALL)
            if match:
                return text[match.end():]
        return text

    # -------------------------------------------------------------------------
    # Style resolution
    # -------------------------------------------------------------------------

    def _get_style(self, key, fallback='Normal'):
        """Get style name from map, validating it exists in the document."""
        style_name = self.style_map.get(key, fallback)
        try:
            self.doc.styles[style_name]
            return style_name
        except KeyError:
            return fallback

    # -------------------------------------------------------------------------
    # Image resolution
    # -------------------------------------------------------------------------

    def resolve_image_path(self, src):
        """Resolve an image src to an absolute path in the vault."""
        if not self.vault_path:
            return None

        # Direct path from vault root
        candidate = self.vault_path / src
        if candidate.exists():
            return candidate

        # Try Attachments/ subfolder
        candidate = self.vault_path / 'Attachments' / src
        if candidate.exists():
            return candidate

        # Basename search in Attachments/
        attachments_dir = self.vault_path / 'Attachments'
        if attachments_dir.exists():
            basename = Path(src).name
            for f in attachments_dir.rglob(basename):
                if f.is_file():
                    return f

        return None

    # -------------------------------------------------------------------------
    # Hyperlink helper (python-docx lacks native API)
    # -------------------------------------------------------------------------

    @staticmethod
    def _add_hyperlink(paragraph, url, display_text):
        """Add a clickable hyperlink to a paragraph using low-level XML."""
        part = paragraph.part
        r_id = part.relate_to(
            url,
            'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
            is_external=True,
        )

        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)

        run_elem = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')

        # Blue underline style
        color = OxmlElement('w:color')
        color.set(qn('w:val'), '0563C1')
        rPr.append(color)
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)

        run_elem.append(rPr)

        text_elem = OxmlElement('w:t')
        text_elem.text = display_text
        text_elem.set(qn('xml:space'), 'preserve')
        run_elem.append(text_elem)

        hyperlink.append(run_elem)
        paragraph._p.append(hyperlink)

    # -------------------------------------------------------------------------
    # Inline rendering (recursive)
    # -------------------------------------------------------------------------

    def _extract_text(self, children):
        """Recursively extract plain text from inline AST nodes."""
        parts = []
        for child in children:
            ctype = child.get('type', '')
            if ctype == 'text':
                parts.append(child.get('raw', child.get('text', '')))
            elif ctype == 'codespan':
                parts.append(child.get('raw', child.get('text', '')))
            elif ctype in ('strong', 'emphasis', 'link'):
                parts.append(self._extract_text(child.get('children', [])))
            elif ctype == 'wikilink':
                parts.append(child.get('attrs', {}).get('alias', ''))
        return ''.join(parts)

    def render_inline(self, paragraph, children, bold=False, italic=False):
        """Recursively render inline AST tokens as runs on a paragraph."""
        for child in children:
            ctype = child.get('type', '')

            if ctype == 'text':
                raw = child.get('raw', child.get('text', ''))
                if raw:
                    run = paragraph.add_run(raw)
                    if bold:
                        run.bold = True
                    if italic:
                        run.italic = True

            elif ctype == 'strong':
                self.render_inline(
                    paragraph, child.get('children', []),
                    bold=True, italic=italic,
                )

            elif ctype == 'emphasis':
                self.render_inline(
                    paragraph, child.get('children', []),
                    bold=bold, italic=True,
                )

            elif ctype == 'codespan':
                raw = child.get('raw', child.get('text', ''))
                run = paragraph.add_run(raw)
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                if bold:
                    run.bold = True
                if italic:
                    run.italic = True

            elif ctype == 'link':
                url = child.get('attrs', {}).get('url', child.get('link', ''))
                display = self._extract_text(child.get('children', []))
                if not display:
                    display = url
                self._add_hyperlink(paragraph, url, display)
                self.stats['hyperlinks'] += 1

            elif ctype == 'image':
                url = child.get('attrs', {}).get('url', child.get('src', ''))
                # Obsidian embeds were preprocessed to obsidian-embed: URLs
                if url.startswith('obsidian-embed:'):
                    src = url[len('obsidian-embed:'):]
                else:
                    src = url
                self._render_image(src)

            elif ctype == 'wikilink':
                alias = child.get('attrs', {}).get('alias', '')
                run = paragraph.add_run(alias)
                if bold:
                    run.bold = True
                if italic:
                    run.italic = True
                self.stats['wikilinks'] += 1

            elif ctype == 'linebreak':
                paragraph.add_run().add_break()

            elif ctype == 'softbreak':
                paragraph.add_run(' ')

            elif ctype == 'strikethrough':
                text = self._extract_text(child.get('children', []))
                run = paragraph.add_run(text)
                run.font.strike = True

    # -------------------------------------------------------------------------
    # Image rendering
    # -------------------------------------------------------------------------

    def _render_image(self, src):
        """Add an image as its own paragraph in the document."""
        image_path = self.resolve_image_path(src)

        if self.dry_run:
            if image_path and image_path.exists():
                self.stats['images']['found'] += 1
            else:
                self.stats['images']['missing'].append(src)
            return

        if image_path and image_path.exists():
            try:
                self.doc.add_picture(str(image_path), width=Inches(5.5))
                self.stats['images']['found'] += 1
            except Exception as e:
                self.doc.add_paragraph(f'[Image error: {src} - {e}]')
                self.stats['images']['missing'].append(src)
        else:
            self.doc.add_paragraph(f'[Image not found: {src}]')
            self.stats['images']['missing'].append(src)

    # -------------------------------------------------------------------------
    # Block-level rendering
    # -------------------------------------------------------------------------

    def render_block(self, token):
        """Dispatch a block-level AST token to the appropriate renderer."""
        ttype = token.get('type', '')

        if ttype == 'heading':
            self._render_heading(token)
        elif ttype == 'paragraph':
            self._render_paragraph(token)
        elif ttype == 'list':
            self._render_list(token)
        elif ttype == 'block_code':
            self._render_code_block(token)
        elif ttype == 'block_quote':
            self._render_block_quote(token)
        elif ttype == 'table':
            self._render_table(token)
        elif ttype == 'thematic_break':
            if not self.dry_run:
                para = self.doc.add_paragraph()
                run = para.add_run()
                run.add_break()
        elif ttype == 'blank_line':
            pass  # ignore
        else:
            # Fallback: render children if any
            children = token.get('children', [])
            if children and isinstance(children, list):
                if isinstance(children[0], dict) and 'type' in children[0]:
                    for child in children:
                        self.render_block(child)

    def _render_heading(self, token):
        level = token.get('attrs', {}).get('level', 1)

        # Title promotion (default): first H1 becomes Title, all others promote by 1
        if not self.no_title_promotion:
            if level == 1 and not self._title_rendered:
                # First H1 → document title (no numbering)
                self._title_rendered = True
                self.stats['headings']['Title'] = 1

                if not self.dry_run:
                    style = self._get_style('title', 'Title')
                    para = self.doc.add_paragraph(style=style)
                    self.render_inline(para, token.get('children', []))
                return

            if self._title_rendered and level > 1:
                # Promote: H2→H1, H3→H2, etc.
                level = level - 1

        style_key = f'heading_{level}'
        h_key = f'H{level}'
        self.stats['headings'][h_key] = self.stats['headings'].get(h_key, 0) + 1

        if not self.dry_run:
            style = self._get_style(style_key, f'Heading {level}')
            para = self.doc.add_paragraph(style=style)
            self.render_inline(para, token.get('children', []))

    def _render_paragraph(self, token):
        children = token.get('children', [])

        # Skip empty paragraphs (no children or only whitespace/softbreaks)
        if not children:
            return
        text_content = self._extract_text(children).strip()
        if not text_content:
            # Check if there's at least an image or embed
            has_content = any(
                c.get('type') in ('image', 'wiki_embed')
                for c in children
            )
            if not has_content:
                return

        self.stats['paragraphs'] += 1

        # Check if paragraph contains only an image
        if len(children) == 1:
            child = children[0]
            ctype = child.get('type', '')
            if ctype == 'image':
                url = child.get('attrs', {}).get('url', child.get('src', ''))
                if url.startswith('obsidian-embed:'):
                    src = url[len('obsidian-embed:'):]
                else:
                    src = url
                self._render_image(src)
                return

        if not self.dry_run:
            style = self._get_style('paragraph')
            para = self.doc.add_paragraph(style=style)
            self.render_inline(para, children)

    def _render_list(self, token, depth=0):
        ordered = token.get('attrs', {}).get('ordered', False)
        if ordered:
            self.stats['lists']['numbered'] += 1
        else:
            self.stats['lists']['bullet'] += 1

        for item in token.get('children', []):
            # item is a list_item
            for child in item.get('children', []):
                ctype = child.get('type', '')
                if ctype in ('paragraph', 'block_text'):
                    if not self.dry_run:
                        if ordered:
                            style_key = f'list_number_{depth}' if depth > 0 else 'list_number'
                        else:
                            style_key = f'list_bullet_{depth}' if depth > 0 else 'list_bullet'
                        style = self._get_style(style_key)
                        para = self.doc.add_paragraph(style=style)
                        self.render_inline(para, child.get('children', []))
                elif ctype == 'list':
                    self._render_list(child, depth=depth + 1)
                else:
                    self.render_block(child)

    def _render_code_block(self, token):
        self.stats['code_blocks'] += 1
        if self.dry_run:
            return

        raw = token.get('raw', token.get('text', ''))
        style = self._get_style('code_block')
        for line in raw.rstrip('\n').split('\n'):
            para = self.doc.add_paragraph(style=style)
            run = para.add_run(line)
            run.font.name = 'Courier New'
            run.font.size = Pt(9)

    @staticmethod
    def _strip_callout_prefix(children):
        """Strip Obsidian callout prefix [!type] from inline children.

        Mistune splits [!type] into tokens like:
          {'type': 'text', 'raw': '['} + {'type': 'text', 'raw': '!info] rest...'}
        Returns (callout_type or None, cleaned children).
        """
        if len(children) < 2:
            return None, children

        first = children[0]
        second = children[1]
        if (first.get('type') == 'text' and first.get('raw', '').strip() == '['
                and second.get('type') == 'text'):
            raw = second.get('raw', '')
            m = re.match(r'^!(\w+)\]\+?\s*', raw)
            if m:
                callout_type = m.group(1)
                remainder = raw[m.end():]
                new_children = []
                if remainder:
                    new_children.append({'type': 'text', 'raw': remainder})
                new_children.extend(children[2:])
                # Strip leading softbreak if the remainder was empty
                if new_children and new_children[0].get('type') == 'softbreak':
                    new_children = new_children[1:]
                return callout_type, new_children
        return None, children

    def _render_block_quote(self, token):
        self.stats['block_quotes'] += 1
        if self.dry_run:
            return

        style = self._get_style('block_quote')
        children = token.get('children', [])

        for child in children:
            ctype = child.get('type', '')
            if ctype == 'paragraph':
                inline_children = child.get('children', [])

                # Strip any remaining callout prefix [!type] (safety net)
                _callout_type, inline_children = self._strip_callout_prefix(inline_children)

                if not inline_children:
                    continue

                para = self.doc.add_paragraph(style=style)
                self.render_inline(para, inline_children)
            elif ctype == 'table':
                # Tables extracted from blockquotes by preprocessing
                self._render_table(child)
            else:
                self.render_block(child)

    def _render_table(self, token):
        self.stats['tables'] += 1
        if self.dry_run:
            return

        children = token.get('children', [])
        if not children:
            return

        # Collect all rows with their cells
        # table_head has table_cell children directly (one header row)
        # table_body has table_row children, each with table_cell children
        all_rows = []  # list of (is_head, [cell_tokens])
        for section in children:
            stype = section.get('type', '')
            if stype == 'table_head':
                # Head cells are direct children
                all_rows.append((True, section.get('children', [])))
            elif stype == 'table_body':
                for row in section.get('children', []):
                    all_rows.append((False, row.get('children', [])))

        if not all_rows:
            return

        num_cols = len(all_rows[0][1])
        num_rows = len(all_rows)

        table = self.doc.add_table(rows=num_rows, cols=num_cols)
        try:
            table.style = self._get_style('table', 'Table Grid')
        except Exception:
            pass

        for row_idx, (is_head, cells) in enumerate(all_rows):
            for col_idx, cell_token in enumerate(cells):
                if col_idx < num_cols:
                    cell = table.rows[row_idx].cells[col_idx]
                    para = cell.paragraphs[0]
                    self.render_inline(
                        para,
                        cell_token.get('children', []),
                        bold=is_head,
                    )

    # -------------------------------------------------------------------------
    # Main execution
    # -------------------------------------------------------------------------

    def parse_note(self):
        """Read and parse the markdown note, returning the AST."""
        text = self.note_path.read_text(encoding='utf-8', errors='ignore')
        text = self.strip_frontmatter(text)
        text = _preprocess_embeds(text)
        text = _preprocess_callouts_and_tables(text)
        return self.md(text), text

    def run(self):
        """Main execution: parse, render, save (or report in dry-run)."""
        ast_tokens, raw_text = self.parse_note()

        if self.dry_run:
            # Walk AST to collect stats without creating DOCX
            for token in ast_tokens:
                self.render_block(token)
            self._print_dry_run_report(raw_text)
            return 0

        # Create document from template
        self.doc = Document(str(self.template_path))

        # Clear any placeholder content in the template
        for para in list(self.doc.paragraphs):
            para._element.getparent().remove(para._element)

        # Render all blocks
        for token in ast_tokens:
            self.render_block(token)

        # Clean-by-default sanitation pass
        if self.sanitize:
            if not self.keep_heading_numbers:
                disable_heading_autonumber(self.doc)
            remove_empty_paragraphs(self.doc)
            normalize_tables(self.doc)

        # Save
        self.doc.save(str(self.output_path))
        print(f'Exported: {self.output_path}')
        total_images = self.stats['images']['found'] + len(self.stats['images']['missing'])
        if total_images:
            missing = len(self.stats['images']['missing'])
            print(f'  Images: {self.stats["images"]["found"]} embedded, {missing} missing')
        return 0

    def _print_dry_run_report(self, raw_text):
        """Print a summary of what would be created."""
        print(f'Obsidian Note to DOCX Exporter')
        print(f'Note:     {self.note_path}')
        print(f'Template: {self.template_path}')
        print(f'Output:   {self.output_path}')
        print(f'Mode:     DRY RUN')
        if not self.no_title_promotion:
            print(f'Title:    First H1 → Title style, remaining headings promoted')
        print('=' * 60)
        print(f'\nContent: {len(raw_text)} characters')
        print(f'\nDocument structure:')

        # Headings
        total_h = sum(self.stats['headings'].values())
        if total_h:
            breakdown = ', '.join(
                f'{k}: {v}' for k, v in sorted(self.stats['headings'].items())
            )
            print(f'  Headings:     {total_h} ({breakdown})')

        print(f'  Paragraphs:   {self.stats["paragraphs"]}')

        total_lists = self.stats['lists']['bullet'] + self.stats['lists']['numbered']
        if total_lists:
            print(f'  Lists:        {total_lists} '
                  f'({self.stats["lists"]["bullet"]} bullet, '
                  f'{self.stats["lists"]["numbered"]} numbered)')

        if self.stats['code_blocks']:
            print(f'  Code blocks:  {self.stats["code_blocks"]}')
        if self.stats['block_quotes']:
            print(f'  Blockquotes:  {self.stats["block_quotes"]}')
        if self.stats['tables']:
            print(f'  Tables:       {self.stats["tables"]}')

        total_img = self.stats['images']['found'] + len(self.stats['images']['missing'])
        if total_img:
            missing = self.stats['images']['missing']
            print(f'  Images:       {total_img} '
                  f'({self.stats["images"]["found"]} found, {len(missing)} missing)')
            for m in missing:
                print(f'                  - Missing: {m}')

        if self.stats['wikilinks']:
            print(f'  Wikilinks:    {self.stats["wikilinks"]} (rendered as plain text)')
        if self.stats['hyperlinks']:
            print(f'  Hyperlinks:   {self.stats["hyperlinks"]}')

        # Validate styles against template
        print(f'\nStyle mapping:')
        try:
            doc = Document(str(self.template_path))
            available = {s.name for s in doc.styles}
            for key, style_name in self.style_map.items():
                status = 'OK' if style_name in available else 'NOT FOUND -> fallback to Normal'
                print(f'  {key:20s} -> "{style_name}" ({status})')
        except Exception as e:
            print(f'  Could not validate styles: {e}')

        print(f'\nThis was a DRY RUN. Run with --no-dry-run to create the DOCX file.')


# =============================================================================
# CLI entry point
# =============================================================================

def main():
    parser = argparse.ArgumentParser(
        description='Export an Obsidian markdown note to DOCX using a template',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Preview what would be exported (dry run)
  python export_to_docx.py --note path/to/note.md --template template.docx

  # Actually export
  python export_to_docx.py --note path/to/note.md --template template.docx --no-dry-run

  # Custom output path
  python export_to_docx.py --note note.md --template t.docx --output result.docx --no-dry-run

  # List styles available in a template
  python export_to_docx.py --template template.docx --list-styles
        """,
    )
    parser.add_argument('--note', type=Path,
                        help='Path to the markdown note to export')
    parser.add_argument('--template', type=Path, required=True,
                        help='Path to .docx template file')
    parser.add_argument('--output', type=Path,
                        help='Output .docx path (default: same name as note)')
    parser.add_argument('--vault', type=str, default='/Users/jose/obsidian/JC',
                        help='Path to Obsidian vault for resolving images')
    parser.add_argument('--no-dry-run', action='store_true',
                        help='Actually create the DOCX (default is dry-run)')
    parser.add_argument('--list-styles', action='store_true',
                        help='List all styles in the template and exit')
    parser.add_argument('--style-map', type=Path,
                        help='JSON file mapping markdown elements to template style names')
    parser.add_argument('--keep-heading-numbers', action='store_true',
                        help='Keep the template\'s heading auto-numbering (default: strip it for prose)')
    parser.add_argument('--no-sanitize', action='store_true',
                        help='Disable the clean-up pass (empty-paragraph removal, table normalisation)')
    parser.add_argument('--no-title-promotion', action='store_true',
                        help='Disable title promotion (default: first H1 becomes Title, '
                             'remaining headings promote one level)')

    args = parser.parse_args()

    # --list-styles mode
    if args.list_styles:
        if not args.template.exists():
            print(f'Error: Template not found: {args.template}')
            return 1
        exporter = MarkdownToDocxExporter(
            note_path=None, template_path=args.template,
            output_path=None, vault_path=args.vault,
            dry_run=True,
        )
        exporter.list_template_styles()
        return 0

    # Normal export mode
    if not args.note:
        parser.error('--note is required unless using --list-styles')

    if not args.note.exists():
        print(f'Error: Note not found: {args.note}')
        return 1
    if not args.template.exists():
        print(f'Error: Template not found: {args.template}')
        return 1

    vault_valid, vault_err = validate_vault_path(args.vault)
    if not vault_valid:
        print(f'Warning: {vault_err} (images may not resolve)')

    output = args.output or args.note.with_suffix('.docx')

    style_overrides = None
    if args.style_map:
        if not args.style_map.exists():
            print(f'Error: Style map not found: {args.style_map}')
            return 1
        with open(args.style_map) as f:
            style_overrides = json.load(f)

    exporter = MarkdownToDocxExporter(
        note_path=args.note,
        template_path=args.template,
        output_path=output,
        vault_path=args.vault,
        dry_run=not args.no_dry_run,
        style_map_overrides=style_overrides,
        no_title_promotion=args.no_title_promotion,
        sanitize=not args.no_sanitize,
        keep_heading_numbers=args.keep_heading_numbers,
    )
    return exporter.run()


if __name__ == '__main__':
    sys.exit(main())
