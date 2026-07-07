#!/usr/bin/env python3
"""
Shared DOCX formatting/sanitation helpers for the obsidian-tools DOCX pipeline
(fill_template.py and export_to_docx.py).

Purpose: make generated TYPSA documents come out clean by default, without a
manual post-processing pass. Fixes the recurring bugs seen in practice:

  - double heading indices (Word auto-numbering + a manual "N." in the text),
  - empty paragraphs and stray empty bullets left by cleared template slots,
  - tables that do not span the text width, with mis-proportioned columns,
    justified cells with large gaps, and rows that split across pages.

All functions operate on a python-docx Document in place.
"""

import re

from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Column-width fractions of the usable text width, keyed by column count.
# Tuned for the two RRAM minutes tables (participants: org/name/function;
# actions: ref/owner/action/timing). Any other column count splits evenly.
DEFAULT_TABLE_FRACTIONS = {
    3: [0.24, 0.30, 0.46],
    4: [0.07, 0.17, 0.56, 0.20],
}


def strip_leading_number(text):
    """Drop a leading manual section number ("1. ", "10.2 ", "3) ") so it does
    not double up with Word's heading auto-numbering."""
    if not text:
        return text
    return re.sub(r'^\s*\d+([.)]\d+)*[.)]\s+', '', text)


def _paragraph_is_empty(p):
    """True only for a genuinely empty paragraph: no text and no drawing,
    picture or break (page breaks and images must be preserved)."""
    if p.text.strip():
        return False
    for tag in ('w:drawing', 'w:pict', 'w:br', 'w:object'):
        if p._p.findall('.//' + qn(tag)):
            return False
    return True


def remove_empty_paragraphs(doc):
    """Delete empty body paragraphs (blank lines and emptied bullet slots)."""
    removed = 0
    for p in list(doc.paragraphs):
        if _paragraph_is_empty(p):
            p._p.getparent().remove(p._p)
            removed += 1
    return removed


def disable_paragraph_numbering(p):
    """Force numId 0 (no list) on a paragraph, overriding any style- or
    direct-numbering it would otherwise inherit."""
    pPr = p._p.get_or_add_pPr()
    for np in pPr.findall(qn('w:numPr')):
        pPr.remove(np)
    numPr = OxmlElement('w:numPr')
    for tag, val in (('w:ilvl', '0'), ('w:numId', '0')):
        e = OxmlElement(tag)
        e.set(qn('w:val'), val)
        numPr.append(e)
    pStyle = pPr.find(qn('w:pStyle'))
    if pStyle is not None:
        pStyle.addnext(numPr)
    else:
        pPr.insert(0, numPr)


def disable_heading_autonumber(doc):
    """Turn off inherited heading/title auto-numbering across the document.
    Use for free-form prose exports where numbered headings are unwanted."""
    for p in doc.paragraphs:
        name = (p.style.name or '') if p.style else ''
        if name.startswith('Heading') or name == 'Title':
            disable_paragraph_numbering(p)


def _usable_twips(section):
    return int((section.page_width - section.left_margin - section.right_margin) / 635)


def _set_table_cell_margins(table, top=28, bottom=28, left=90, right=90):
    tblPr = table._tbl.tblPr
    old = tblPr.find(qn('w:tblCellMar'))
    if old is not None:
        tblPr.remove(old)
    mar = OxmlElement('w:tblCellMar')
    for side, val in (('top', top), ('bottom', bottom), ('left', left), ('right', right)):
        e = OxmlElement('w:' + side)
        e.set(qn('w:w'), str(val))
        e.set(qn('w:type'), 'dxa')
        mar.append(e)
    tblPr.append(mar)


def _set_fixed_layout(table, total):
    tblPr = table._tbl.tblPr
    for tag in ('w:tblW', 'w:tblLayout'):
        e = tblPr.find(qn(tag))
        if e is not None:
            tblPr.remove(e)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(total))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)
    lay = OxmlElement('w:tblLayout')
    lay.set(qn('w:type'), 'fixed')
    tblPr.append(lay)


def normalize_tables(doc, fractions=None):
    """Give every table a full-width fixed layout with proportional columns,
    left-aligned cells, tight cell margins, and rows that do not split across
    pages. Removes empty leading/trailing paragraphs inside cells."""
    fractions = fractions or DEFAULT_TABLE_FRACTIONS
    usable = _usable_twips(doc.sections[0])

    for t in doc.tables:
        ncols = len(t.columns)
        fr = fractions.get(ncols, [1.0 / ncols] * ncols)
        widths = [int(f * usable) for f in fr]
        widths[-1] = usable - sum(widths[:-1])  # keep the exact total

        _set_fixed_layout(t, usable)
        _set_table_cell_margins(t)

        grid = t._tbl.find(qn('w:tblGrid'))
        if grid is not None:
            for c in grid.findall(qn('w:gridCol')):
                grid.remove(c)
            for w in widths:
                gc = OxmlElement('w:gridCol')
                gc.set(qn('w:w'), str(w))
                grid.append(gc)

        for row in t.rows:
            trPr = row._tr.get_or_add_trPr()
            if trPr.find(qn('w:cantSplit')) is None:
                trPr.append(OxmlElement('w:cantSplit'))
            for i, cell in enumerate(row.cells):
                tcPr = cell._tc.get_or_add_tcPr()
                tcW = tcPr.find(qn('w:tcW'))
                if tcW is None:
                    tcW = OxmlElement('w:tcW')
                    tcPr.append(tcW)
                tcW.set(qn('w:w'), str(widths[i] if i < len(widths) else widths[-1]))
                tcW.set(qn('w:type'), 'dxa')
                cps = cell.paragraphs
                for p in cps:
                    if len(cps) > 1 and _paragraph_is_empty(p):
                        p._p.getparent().remove(p._p)
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    pf = p.paragraph_format
                    pf.space_before = Pt(0)
                    pf.space_after = Pt(1)
