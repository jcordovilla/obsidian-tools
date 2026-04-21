#!/usr/bin/env python3
"""
DOCX Template Filler

Fills structured DOCX templates with data from YAML files using registered profiles.
Each profile maps a YAML data schema to specific template elements (paragraphs, tables).

Usage:
    python fill_template.py --template t.docx --data data.yaml --profile rram-minutes
    python fill_template.py --template t.docx --data data.yaml --profile rram-minutes --output out.docx --no-dry-run
    python fill_template.py --list-profiles
"""

import argparse
import copy
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("Error: python-docx not installed. Install with: pip install python-docx")
    sys.exit(1)

try:
    import yaml
except ImportError:
    print("Error: pyyaml not installed. Install with: pip install pyyaml")
    sys.exit(1)


# =============================================================================
# Low-level helpers
# =============================================================================

def _set_paragraph_text(para, text):
    """Replace paragraph text, preserving style and first run's character formatting."""
    runs = list(para._p.findall(qn('w:r')))
    if runs and text:
        first_r = runs[0]
        t = first_r.find(qn('w:t'))
        if t is None:
            t = OxmlElement('w:t')
            first_r.append(t)
        t.text = text
        t.set(qn('xml:space'), 'preserve')
        for r in runs[1:]:
            para._p.remove(r)
    elif text:
        para.add_run(text)
    else:
        for r in runs:
            para._p.remove(r)


def _remove_paragraph(para):
    """Remove a paragraph from the document body."""
    p = para._p
    parent = p.getparent()
    if parent is not None:
        parent.remove(p)


def _set_cell_text(cell, text):
    """Set table cell text, preserving first run's character formatting."""
    para = cell.paragraphs[0]
    runs = list(para._p.findall(qn('w:r')))
    if runs:
        first_r = runs[0]
        t = first_r.find(qn('w:t'))
        if t is None:
            t = OxmlElement('w:t')
            first_r.append(t)
        t.text = text
        t.set(qn('xml:space'), 'preserve')
        for r in runs[1:]:
            para._p.remove(r)
    else:
        para.add_run(text)


def _fill_table(table, rows, header_rows=1, placeholder_rows=3):
    """Replace placeholder data rows with actual data rows, preserving style."""
    # Deep-copy first data row as style template
    style_tr = None
    if len(table.rows) > header_rows:
        style_tr = copy.deepcopy(table.rows[header_rows]._tr)

    # Collect and remove placeholder data rows
    trs = []
    for i in range(header_rows, min(header_rows + placeholder_rows, len(table.rows))):
        trs.append(table.rows[i]._tr)
    for tr in trs:
        table._tbl.remove(tr)

    # Add new data rows
    for row_data in rows:
        if style_tr is not None:
            new_tr = copy.deepcopy(style_tr)
            table._tbl.append(new_tr)
            row = table.rows[-1]
            for col_idx, value in enumerate(row_data):
                if col_idx < len(row.cells):
                    _set_cell_text(row.cells[col_idx], str(value))
        else:
            row = table.add_row()
            for col_idx, value in enumerate(row_data):
                if col_idx < len(row.cells):
                    row.cells[col_idx].text = str(value)


# =============================================================================
# RRAM Minutes profile
# =============================================================================

# Template paragraph indices (0-based) and table indices
_RRAM = {
    'subtitle': 1,
    'participants_table': 0,
    'participants_placeholder_rows': 3,
    'agenda': [
        {'heading': 3,  'body': 4},
        {'heading': 5,  'body': 6},
        {'heading': 7,  'body': 8},
        {'heading': 9,  'body': 10, 'extra': [11, 12, 13, 14, 15]},
        {'heading': 16, 'body': 17},
        {'heading': 18, 'body': 19},
        {'heading': 20, 'body': 21},
    ],
    'next_steps_intro': 23,
    'next_steps': [24, 25, 26],
    'actions_table': 1,
    'actions_placeholder_rows': 3,
}


def _validate_rram(data):
    """Validate RRAM minutes YAML data."""
    errors = []
    for field in ('subtitle', 'participants', 'agenda'):
        if field not in data:
            errors.append(f"Missing required field: {field}")
    if 'participants' in data:
        if not isinstance(data['participants'], list):
            errors.append("'participants' must be a list")
        else:
            for i, p in enumerate(data['participants']):
                for f in ('org', 'name', 'function'):
                    if f not in p:
                        errors.append(f"participants[{i}] missing '{f}'")
    if 'agenda' in data:
        if not isinstance(data['agenda'], list):
            errors.append("'agenda' must be a list")
        elif len(data['agenda']) > 7:
            errors.append(f"Template supports max 7 agenda items, got {len(data['agenda'])}")
    if 'actions' in data:
        for i, a in enumerate(data['actions']):
            for f in ('ref', 'owner', 'action', 'timing'):
                if f not in a:
                    errors.append(f"actions[{i}] missing '{f}'")
    return errors


def _fill_rram(doc, data, dry_run=True):
    """Fill RRAM meeting minutes template with data."""
    paras = doc.paragraphs
    tables = doc.tables
    report = []
    to_remove = []

    # 1. Subtitle
    subtitle = data['subtitle']
    report.append(f"Subtitle: {subtitle}")
    if not dry_run:
        _set_paragraph_text(paras[_RRAM['subtitle']], subtitle)

    # 2. Participants table
    participants = data['participants']
    report.append(f"\nParticipants ({len(participants)}):")
    for p in participants:
        report.append(f"  {p['org']:20s} | {p['name']:25s} | {p['function']}")
    if not dry_run:
        rows = [[p['org'], p['name'], p['function']] for p in participants]
        _fill_table(tables[_RRAM['participants_table']], rows,
                    placeholder_rows=_RRAM['participants_placeholder_rows'])

    # 3. Agenda items
    agenda = data['agenda']
    report.append(f"\nAgenda ({len(agenda)} items):")
    for slot_idx, slot in enumerate(_RRAM['agenda']):
        if slot_idx < len(agenda):
            item = agenda[slot_idx]
            title = item['title']
            body_parts = item.get('body', [])
            if isinstance(body_parts, str):
                body_parts = [body_parts]
            body = ' '.join(body_parts)

            report.append(f"  {title}")
            report.append(f"    {body[:80]}{'...' if len(body) > 80 else ''}")

            if not dry_run:
                _set_paragraph_text(paras[slot['heading']], title)
                _set_paragraph_text(paras[slot['body']], body)
                # Clear extra paragraph slots (bullet placeholders)
                for idx in slot.get('extra', []):
                    _set_paragraph_text(paras[idx], '')
        else:
            report.append(f"  (slot {slot_idx + 1} unused)")
            to_remove.append(paras[slot['heading']])
            to_remove.append(paras[slot['body']])
            for idx in slot.get('extra', []):
                to_remove.append(paras[idx])

    # 4. Next steps
    next_steps = data.get('next_steps', [])
    report.append(f"\nNext steps ({len(next_steps)}):")
    for ns in next_steps:
        report.append(f"  - {ns[:80]}{'...' if len(ns) > 80 else ''}")
    if not dry_run:
        for slot_idx, p_idx in enumerate(_RRAM['next_steps']):
            if slot_idx < len(next_steps):
                _set_paragraph_text(paras[p_idx], next_steps[slot_idx])
            else:
                _set_paragraph_text(paras[p_idx], '')

    # 5. Actions table
    actions = data.get('actions', [])
    report.append(f"\nActions ({len(actions)}):")
    for a in actions:
        report.append(f"  [{a['ref']}] {a['owner']}: {a['action'][:60]}...")
        report.append(f"       Timing: {a['timing']}")
    if not dry_run:
        rows = [[a['ref'], a['owner'], a['action'], a['timing']] for a in actions]
        _fill_table(tables[_RRAM['actions_table']], rows,
                    placeholder_rows=_RRAM['actions_placeholder_rows'])

    # 6. Remove unused paragraphs (after all fill operations)
    if not dry_run:
        for para in to_remove:
            _remove_paragraph(para)

    return report


# =============================================================================
# Profile registry
# =============================================================================

PROFILES = {
    'rram-minutes': {
        'name': 'RRAM Meeting Minutes',
        'description': 'SSATP RRAM project (IA2347) meeting minutes',
        'validate': _validate_rram,
        'fill': _fill_rram,
    },
}


# =============================================================================
# CLI
# =============================================================================

def main():
    parser = argparse.ArgumentParser(
        description='Fill DOCX templates with YAML data using profiles',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Preview (dry run)
  python fill_template.py --template t.docx --data data.yaml --profile rram-minutes

  # Fill template
  python fill_template.py --template t.docx --data d.yaml --profile rram-minutes \\
      --output out.docx --no-dry-run

  # List profiles
  python fill_template.py --list-profiles
        """,
    )
    parser.add_argument('--template', type=Path,
                        help='Path to .docx template')
    parser.add_argument('--data', type=Path,
                        help='Path to YAML data file')
    parser.add_argument('--profile', type=str,
                        help='Profile name (see --list-profiles)')
    parser.add_argument('--output', type=Path,
                        help='Output .docx path (default: <template>-filled.docx)')
    parser.add_argument('--no-dry-run', action='store_true',
                        help='Actually fill the template (default is dry-run)')
    parser.add_argument('--list-profiles', action='store_true',
                        help='List available profiles and exit')

    args = parser.parse_args()

    if args.list_profiles:
        print("Available profiles:")
        for name, info in PROFILES.items():
            print(f"  {name:20s}  {info['description']}")
        return 0

    if not args.template or not args.data or not args.profile:
        parser.error("--template, --data, and --profile are required")
    if not args.template.exists():
        print(f"Error: Template not found: {args.template}")
        return 1
    if not args.data.exists():
        print(f"Error: Data file not found: {args.data}")
        return 1
    if args.profile not in PROFILES:
        print(f"Error: Unknown profile '{args.profile}'. Use --list-profiles")
        return 1

    profile = PROFILES[args.profile]

    with open(args.data) as f:
        data = yaml.safe_load(f)

    errors = profile['validate'](data)
    if errors:
        print("Validation errors:")
        for e in errors:
            print(f"  - {e}")
        return 1

    output = args.output or args.template.with_name(
        args.template.stem + '-filled' + args.template.suffix
    )

    doc = Document(str(args.template))
    dry_run = not args.no_dry_run
    report = profile['fill'](doc, data, dry_run=dry_run)

    header = "DRY RUN" if dry_run else "FILLING"
    print(f"Template Filler -- {header}")
    print(f"  Template: {args.template}")
    print(f"  Data:     {args.data}")
    print(f"  Profile:  {args.profile} ({profile['name']})")
    print(f"  Output:   {output}")
    print("=" * 60)
    for line in report:
        print(line)

    if not dry_run:
        doc.save(str(output))
        print(f"\nSaved: {output}")
    else:
        print(f"\nDry run complete. Use --no-dry-run to create the document.")

    return 0


if __name__ == '__main__':
    sys.exit(main())
